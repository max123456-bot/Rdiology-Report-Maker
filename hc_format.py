"""
HC FORMAT engine for radiology reports.

Two stages, deliberately separated:

    parse_report(raw_text) -> list[Block]     # structure detection only
    build_docx(blocks, ...) -> bytes          # styling only

Nothing in this module rewrites, summarises or reflows the clinical text.
Every character that carries meaning survives from input to output; the only
transformations applied are (a) UPPERCASE on the title / main headings, which
the HC FORMAT demands, and (b) removal of literal bullet glyphs ("-", "*",
"1.") that are replaced by real Word list bullets.  Both are tracked so
verify.py can prove no word was lost.

HC FORMAT rules implemented here
--------------------------------
1. Arial, 12 pt, black only.
2. Normal (1 inch) margins, professional spacing.
3. Report title (first line): centred, BOLD, UNDERLINED, UPPERCASE.
4. Main headings: left aligned, BOLD, UNDERLINED, UPPERCASE.
5. Findings: every finding a bullet; organ subheadings italic + underlined,
   NOT bold.
6. Impression: every point a bullet, every bullet BOLD.
7. Content preserved verbatim.
8. Delivered as .docx.
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

FONT_NAME = "Arial"
BLACK = RGBColor(0, 0, 0)
# Rule 2: professional spacing throughout the document.
LINE_SPACING = 1.5

# Main headings, longest first so "CLINICAL HISTORY" wins over "CLINICAL".
#
# Deliberately absent: bare "PATIENT". Reports that open with "Patient: <clinical
# context>" are giving background, not a patient name, so that line stays plain
# text. Reviewed and approved that way - do not add "PATIENT" here.
MAIN_HEADINGS: tuple[str, ...] = (
    "PATIENT NAME",
    "PATIENT'S NAME",
    "NAME OF PATIENT",
    "AGE/SEX",
    "AGE / SEX",
    "AGE AND SEX",
    "AGE",
    "SEX",
    "REFERRED BY",
    "REFERRING DOCTOR",
    "DATE OF EXAMINATION",
    "DATE",
    "EXAMINATION",
    "STUDY",
    "PROCEDURE",
    "CLINICAL HISTORY",
    "CLINICAL INDICATION",
    "CLINICAL DETAILS",
    "HISTORY",
    "INDICATION",
    "TECHNIQUE",
    "PROTOCOL",
    "IMAGING SEQUENCES USED",
    "SEQUENCES",
    "COMPARISON",
    "CONTRAST",
    "FINDINGS",
    "OBSERVATIONS",
    "IMPRESSION",
    "CONCLUSION",
    "OPINION",
    "COMMENT",
    "COMMENTS",
    "RECOMMENDATION",
    "RECOMMENDATIONS",
    "ADVICE",
    "NOTE",
    "LIMITATIONS",
    "DISCLAIMER",
)

# Short metadata headings that read naturally with their value on the same line
# ("PATIENT NAME: Mrs. Sunita Devi"). Only the heading run is bold+underlined,
# so rule 4 still holds. Long prose sections always get their own paragraph.
INLINE_META_HEADINGS = {
    "PATIENT NAME",
    "PATIENT'S NAME",
    "NAME OF PATIENT",
    "AGE/SEX",
    "AGE / SEX",
    "AGE AND SEX",
    "AGE",
    "SEX",
    "REFERRED BY",
    "REFERRING DOCTOR",
    "DATE OF EXAMINATION",
    "DATE",
    "EXAMINATION",
    "STUDY",
    "PROCEDURE",
}

FINDINGS_SECTIONS = {"FINDINGS", "OBSERVATIONS"}
IMPRESSION_SECTIONS = {"IMPRESSION", "CONCLUSION", "OPINION"}
# Sections whose bullets follow the IMPRESSION rule (bold) per HC FORMAT
# only apply to IMPRESSION/CONCLUSION; COMMENT and RECOMMENDATION stay plain
# bullets unless the user turns bold_comment on.
BULLETED_TAIL_SECTIONS = {"COMMENT", "COMMENTS", "RECOMMENDATION", "RECOMMENDATIONS", "ADVICE"}

# Leading list markers we strip because Word supplies a real bullet instead.
_BULLET_PREFIX = re.compile(r"^\s*(?:[•●▪‣⁃·*–—-]+|\(?\d{1,2}[.)]|\(?[a-z][.)])\s+")
# A colon-terminated short label, optionally with text after it: "Liver:" / "Liver: normal."
_LABEL_SPLIT = re.compile(r"^([^:\n]{1,60}?)\s*:\s*(.*)$", re.DOTALL)
_SENTENCE_SPLIT = re.compile(r"(?<=[.;])\s+(?=[A-Z(])")

# Lines that close a report and must never become bullets.
_SIGNOFF_HINTS = (
    "please correlate",
    "kindly correlate",
    "clinical correlation",
    "correlate clinically",
    "thanks for the referral",
    "thank you for the referral",
    "suggested",
    "dr.",
    "dr ",
    "m.d.",
    "md ",
    "dmrd",
    "dnb",
    "consultant radiologist",
    "radiologist",
    "not a legal document",
    "this report",
    "disclaimer",
)


@dataclass
class Span:
    """A stretch of text carrying its own emphasis, from the rich-text editor.

    The flags are *additions* on top of whatever the line's template style
    already applies. Italicising a word inside an IMPRESSION bullet gives
    bold + italic; the rule-6 bold is never lost.
    """

    text: str
    bold: bool = False
    italic: bool = False
    underline: bool = False


@dataclass
class Block:
    """One rendered paragraph of the output document."""

    kind: str  # title | heading | heading_inline | subheading | bullet | bold_bullet | text | verbatim | spacer
    text: str = ""
    trailer: str = ""  # text that followed an inline label, e.g. "Liver: <trailer>"
    section: str = ""  # owning main heading, for debugging / preview
    dropped_marker: str = ""  # bullet glyph removed, kept for the audit trail
    raw: str = ""  # original source line, byte-for-byte
    spans: list[Span] = field(default_factory=list)  # empty -> render `text` as one run
    trailer_spans: list[Span] = field(default_factory=list)


@dataclass
class ParseOptions:
    split_sentences: bool = False  # split a findings paragraph into one bullet per sentence
    bold_comment_bullets: bool = False  # apply the IMPRESSION bold rule to COMMENT/RECOMMENDATION
    max_subheading_len: int = 45  # a findings line shorter than this with no full stop is a subheading
    inline_meta_headings: bool = True  # "PATIENT NAME: value" on one line, heading run only bold+underlined
    preserve_as_is: bool = False  # print the input exactly: no bullets, no uppercase, spacing kept


@dataclass
class ParseResult:
    blocks: list[Block] = field(default_factory=list)
    title: str = ""
    sections_found: list[str] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)


# --------------------------------------------------------------------------- #
# Parsing
# --------------------------------------------------------------------------- #


def _match_main_heading(line: str) -> tuple[str, str, str] | None:
    """
    Return (canonical_heading, heading_as_written, inline_content) when `line`
    opens a main heading, else None.

    Handles all three real-world shapes:
        FINDINGS:
        FINDINGS
        CLINICAL HISTORY: 45 year old male with headache
    """
    probe = line.strip().lstrip("•*-–— ").strip()
    upper = probe.upper()
    for heading in sorted(MAIN_HEADINGS, key=len, reverse=True):
        if not upper.startswith(heading):
            continue
        rest = probe[len(heading):]
        stripped = rest.lstrip()
        # Must be followed by end-of-line, a colon, or a dash - otherwise
        # "FINDINGS ARE UNREMARKABLE" would be mistaken for a heading.
        if stripped[:1] in ("", ":", "-", "–", "—"):
            inline = stripped[1:].strip() if stripped[:1] else ""
            written = probe[: len(heading)] + (":" if stripped[:1] == ":" else "")
            return heading, written, inline
    return None


def _looks_like_signoff(line: str) -> bool:
    low = line.strip().lower()
    return any(hint in low for hint in _SIGNOFF_HINTS)


def _looks_like_subheading(line: str, opts: ParseOptions) -> bool:
    """Organ / region label inside FINDINGS, e.g. 'Liver', 'Cervical Spine:'."""
    text = line.strip()
    if not text:
        return False
    if text.endswith(":"):
        return True
    if len(text) > opts.max_subheading_len:
        return False
    if text.endswith((".", ",", ";")):
        return False
    # Reject anything that reads as a sentence: measurements, verbs, etc.
    if re.search(r"\d", text) and not re.match(r"^[A-Za-z ]+ ?\d*$", text):
        return False
    words = text.split()
    if len(words) > 5:
        return False
    return True


def parse_report(raw_text: str, opts: ParseOptions | None = None) -> ParseResult:
    """Turn raw report text into ordered Blocks. Never edits clinical wording."""
    opts = opts or ParseOptions()
    result = ParseResult()

    normalised = raw_text.replace("\r\n", "\n").replace("\r", "\n")

    if opts.preserve_as_is:
        return _parse_verbatim(normalised, result)

    lines = [ln.rstrip() for ln in normalised.split("\n")]
    lines = [ln for ln in lines if ln.strip()]
    if not lines:
        result.warnings.append("Input is empty.")
        return result

    # Rule 3: the first line is the report title, whatever it says.
    title_raw = lines[0].strip()
    result.title = title_raw
    result.blocks.append(Block(kind="title", text=title_raw.upper(), raw=title_raw))

    current = "GENERAL"
    for raw_line in lines[1:]:
        line = raw_line.strip()

        heading = _match_main_heading(line)
        if heading:
            canonical, written, inline = heading
            current = canonical
            if canonical not in result.sections_found:
                result.sections_found.append(canonical)

            if inline and opts.inline_meta_headings and canonical in INLINE_META_HEADINGS:
                result.blocks.append(
                    Block(
                        kind="heading_inline",
                        text=written.upper(),
                        trailer=inline,
                        section=canonical,
                        raw=raw_line,
                    )
                )
                continue

            result.blocks.append(
                Block(kind="heading", text=written.upper(), section=canonical, raw=raw_line)
            )
            if inline:
                result.blocks.extend(
                    _content_blocks(inline, current, opts, raw_line, dropped="")
                )
            continue

        marker = ""
        m = _BULLET_PREFIX.match(line)
        if m:
            marker = m.group(0).strip()
            line = line[m.end():].strip()
        if not line:
            continue

        result.blocks.extend(_content_blocks(line, current, opts, raw_line, dropped=marker))

    if "FINDINGS" not in result.sections_found and "OBSERVATIONS" not in result.sections_found:
        result.warnings.append(
            "No FINDINGS/OBSERVATIONS heading detected - those lines were written as plain text."
        )
    if not (set(result.sections_found) & IMPRESSION_SECTIONS):
        result.warnings.append("No IMPRESSION/CONCLUSION heading detected.")
    return result


def _parse_verbatim(text: str, result: ParseResult) -> ParseResult:
    """
    As-is mode: print exactly what was pasted.

    One source line becomes one paragraph, blank lines included. Nothing is
    upper-cased, no list marker is stripped, no bullet is added and leading
    indentation is kept - so text copied out of another system lands on the
    page character for character. Word's `xml:space="preserve"` keeps the
    leading spaces, and tabs survive as real tab characters.
    """
    lines = text.split("\n")
    while lines and not lines[-1].strip():
        lines.pop()  # trailing blank lines from the paste, not from the report

    if not lines:
        result.warnings.append("Input is empty.")
        return result

    result.title = lines[0].strip()
    for line in lines:
        result.blocks.append(Block(kind="verbatim", text=line.rstrip("\n"), raw=line))
    result.warnings.append(
        "As-is mode: the report is printed exactly as pasted. HC FORMAT rules 3-6 "
        "(title, headings, bullets, bold impression) are not applied."
    )
    return result


def _content_blocks(
    line: str, section: str, opts: ParseOptions, raw_line: str, dropped: str
) -> list[Block]:
    """Render one non-heading source line into one or more Blocks."""
    if section in FINDINGS_SECTIONS:
        if _looks_like_signoff(line) and not _looks_like_subheading(line, opts):
            return [Block(kind="text", text=line, section=section, raw=raw_line,
                          dropped_marker=dropped)]

        if _looks_like_subheading(line, opts):
            return [Block(kind="subheading", text=line, section=section, raw=raw_line,
                          dropped_marker=dropped)]

        # "Liver: normal in size and outline." -> italic-underlined label + bullet body.
        m = _LABEL_SPLIT.match(line)
        if m and _looks_like_subheading(m.group(1) + ":", opts) and m.group(2).strip():
            label, body = m.group(1).strip(), m.group(2).strip()
            blocks = [Block(kind="subheading", text=label + ":", section=section, raw=raw_line)]
            blocks.extend(
                Block(kind="bullet", text=part, section=section, raw=raw_line)
                for part in _split_body(body, opts)
            )
            return blocks

        return [
            Block(kind="bullet", text=part, section=section, raw=raw_line,
                  dropped_marker=dropped if i == 0 else "")
            for i, part in enumerate(_split_body(line, opts))
        ]

    if section in IMPRESSION_SECTIONS or (
        opts.bold_comment_bullets and section in BULLETED_TAIL_SECTIONS
    ):
        if _looks_like_signoff(line):
            return [Block(kind="text", text=line, section=section, raw=raw_line,
                          dropped_marker=dropped)]
        return [
            Block(kind="bold_bullet", text=part, section=section, raw=raw_line,
                  dropped_marker=dropped if i == 0 else "")
            for i, part in enumerate(_split_body(line, opts))
        ]

    if section in BULLETED_TAIL_SECTIONS:
        if _looks_like_signoff(line):
            return [Block(kind="text", text=line, section=section, raw=raw_line,
                          dropped_marker=dropped)]
        return [Block(kind="bullet", text=line, section=section, raw=raw_line,
                      dropped_marker=dropped)]

    return [Block(kind="text", text=line, section=section, raw=raw_line, dropped_marker=dropped)]


def _split_body(text: str, opts: ParseOptions) -> list[str]:
    if not opts.split_sentences:
        return [text]
    parts = [p.strip() for p in _SENTENCE_SPLIT.split(text) if p.strip()]
    return parts or [text]


# --------------------------------------------------------------------------- #
# DOCX rendering
# --------------------------------------------------------------------------- #


_ALIGN = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def _rgb(hex_colour: str) -> RGBColor:
    try:
        return RGBColor.from_string((hex_colour or "000000").lstrip("#").upper())
    except Exception:
        return BLACK


# XML 1.0 permits tab, newline and carriage return, and nothing else below
# U+0020. Text copied out of a PACS, a RIS export or a PDF regularly carries
# stray control bytes, and python-docx refuses the whole document over one of
# them - so a single invisible character would take the app down mid-clinic.
_ILLEGAL_XML = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f￾￿]")


def sanitise(text: str) -> str:
    """Strip characters Word cannot store. Removes nothing a reader can see."""
    return _ILLEGAL_XML.sub("", text or "")


def _style_run(
    run,
    *,
    bold=False,
    italic=False,
    underline=False,
    font_name: str = FONT_NAME,
    font_size: float = 12.0,
    colour: RGBColor | None = None,
) -> None:
    font = run.font
    font.name = font_name
    font.size = Pt(font_size)
    font.color.rgb = colour if colour is not None else BLACK
    font.bold = bold
    font.italic = italic
    font.underline = underline
    # python-docx only sets w:ascii; Word falls back to Calibri for other
    # ranges (degree signs, x-signs in "3 x 4 mm") without these.
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), font_name)


def _base_document(template) -> Document:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(template.margin_top)
        section.bottom_margin = Inches(template.margin_bottom)
        section.left_margin = Inches(template.margin_left)
        section.right_margin = Inches(template.margin_right)

    # Applied at style level so bullets and any stray runs inherit it.
    for style_name in ("Normal", "List Bullet"):
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue
        style.font.name = template.font_name
        style.font.size = Pt(template.font_size)
        style.font.color.rgb = _rgb(template.font_color)
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = rpr.makeelement(qn("w:rFonts"), {})
            rpr.append(rfonts)
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rfonts.set(qn(attr), template.font_name)
        pf = style.paragraph_format
        pf.space_after = Pt(4)
        pf.line_spacing = template.line_spacing

    return doc


def _add_letterhead(doc: Document, letterhead: dict, template) -> None:
    """Optional clinic header: logo image, name, address, contact."""
    logo_bytes = letterhead.get("logo_bytes")
    if logo_bytes:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        try:
            p.add_run().add_picture(io.BytesIO(logo_bytes), width=Inches(letterhead.get("logo_width_in", 1.6)))
        except Exception:  # unreadable image must not kill the report
            pass

    for key, bold, size in (("name", True, 14), ("address", False, 10), ("contact", False, 10)):
        value = (letterhead.get(key) or "").strip()
        if not value:
            continue
        for line in value.split("\n"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            # The letterhead is a tight header block, not report body: the report's
            # 1.5 spacing would leave a three-line address looking adrift. Set it
            # explicitly rather than inheriting whatever the body uses.
            p.paragraph_format.line_spacing = 1.0
            _style_run(
                p.add_run(sanitise(line)),
                bold=bold,
                font_name=template.font_name,
                font_size=size,
                colour=_rgb(template.font_color),
            )

    # Horizontal rule under the letterhead.
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    pbdr = p._p.get_or_add_pPr().makeelement(qn("w:pBdr"), {})
    bottom = pbdr.makeelement(qn("w:bottom"), {})
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pbdr.append(bottom)
    p._p.get_or_add_pPr().append(pbdr)


def build_docx(
    blocks: Iterable[Block],
    *,
    template=None,
    letterhead: dict | None = None,
    page_numbers: bool = False,
) -> bytes:
    """
    Render Blocks into a .docx and return the file bytes.

    `template` decides the whole look - font, size, colour, line spacing,
    margins, and the bold/italic/underline/alignment/bullet of each kind of
    line. Omit it and you get the signed-off HC FORMAT.
    """
    from templates import HC_FORMAT

    tpl = template or HC_FORMAT
    doc = _base_document(tpl)
    colour = _rgb(tpl.font_color)

    if letterhead and any(letterhead.get(k) for k in ("name", "address", "contact", "logo_bytes")):
        _add_letterhead(doc, letterhead, tpl)

    def paragraph(style_spec, *, listed: bool):
        """A paragraph carrying the template's spacing and alignment.

        Line spacing is written on the paragraph as well as on the Word style,
        so the value survives someone re-applying a style over the text later.
        """
        p = doc.add_paragraph(style="List Bullet") if listed else doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing = tpl.line_spacing
        pf.space_before = Pt(style_spec.space_before)
        pf.space_after = Pt(style_spec.space_after)
        p.alignment = _ALIGN.get(style_spec.align, WD_ALIGN_PARAGRAPH.LEFT)
        return p

    def write(p, text: str, spans: list[Span], spec, *, force_plain: bool = False) -> None:
        """One run per span, each merging the line's style with the span's own."""
        base_bold = False if force_plain else spec.bold
        base_italic = False if force_plain else spec.italic
        base_underline = False if force_plain else spec.underline
        parts = spans or [Span(text=text)]
        for span in parts:
            clean = sanitise(span.text)
            if clean == "":
                continue
            _style_run(
                p.add_run(clean),
                bold=base_bold or span.bold,
                italic=base_italic or span.italic,
                underline=base_underline or span.underline,
                font_name=spec.font_name or tpl.font_name,
                font_size=spec.font_size or tpl.font_size,
                colour=colour,
            )

    for block in blocks:
        if block.kind == "spacer":
            paragraph(tpl.style("text"), listed=False)
            continue

        spec = tpl.style(block.kind)
        text = block.text.upper() if spec.uppercase else block.text
        p = paragraph(spec, listed=spec.bullet)
        write(p, text, block.spans, spec)

        if block.kind == "heading_inline" and block.trailer:
            # The value keeps the line's font but none of the heading emphasis.
            write(p, " " + block.trailer, _prefixed(block.trailer_spans), spec,
                  force_plain=True)

    if page_numbers:
        _add_page_numbers(doc, tpl)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _prefixed(spans: list[Span]) -> list[Span]:
    """Re-attach the separating space when an inline value carries spans."""
    if not spans:
        return []
    first = spans[0]
    return [Span(" " + first.text, first.bold, first.italic, first.underline)] + spans[1:]


def _add_page_numbers(doc: Document, tpl) -> None:
    for section in doc.sections:
        p = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Page ")
        _style_run(run, font_name=tpl.font_name, font_size=9)
        for instr in ("PAGE", None, "NUMPAGES"):
            if instr is None:
                sep = p.add_run(" of ")
                _style_run(sep, font_name=tpl.font_name, font_size=9)
                continue
            fld = p.add_run()
            _style_run(fld, font_name=tpl.font_name, font_size=9)
            begin = fld._element.makeelement(qn("w:fldChar"), {})
            begin.set(qn("w:fldCharType"), "begin")
            text = fld._element.makeelement(qn("w:instrText"), {})
            text.set(qn("xml:space"), "preserve")
            text.text = f" {instr} "
            end = fld._element.makeelement(qn("w:fldChar"), {})
            end.set(qn("w:fldCharType"), "end")
            fld._element.append(begin)
            fld._element.append(text)
            fld._element.append(end)
