"""
Input readers: get raw report text out of whatever the boss sent.

.txt / .md   - decoded directly
.docx        - paragraphs and table cells, in order
.pdf         - embedded text layer via pypdf; if the PDF is a scan there is no
               text layer, and the caller is told to use the AI OCR path
.png / .jpg  - no local OCR; handled by ai_parser.extract_text_from_file
"""

from __future__ import annotations

import io

TEXT_EXT = {".txt", ".md", ".text"}
IMAGE_EXT = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif", ".tiff"}

MIME_BY_EXT = {
    ".png": "image/png",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".webp": "image/webp",
    ".bmp": "image/bmp",
    ".tif": "image/tiff",
    ".tiff": "image/tiff",
    ".pdf": "application/pdf",
}


class NeedsOCR(Exception):
    """The file carries no extractable text layer - AI vision is required."""


def read_text_file(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


class UnreadableFile(Exception):
    """The file could not be parsed. Message is safe to show a user."""


# A radiology report is text. Anything past this is not a report we should be
# feeding to a parser, and refusing early avoids a zip bomb or a malformed PDF
# taking the app down.
MAX_BYTES = 25 * 1024 * 1024


def _guard_size(data: bytes) -> None:
    if len(data) > MAX_BYTES:
        raise UnreadableFile(
            f"That file is {len(data) // (1024 * 1024)} MB. The limit is "
            f"{MAX_BYTES // (1024 * 1024)} MB — a report should be far smaller than this."
        )


def read_docx(data: bytes) -> str:
    from docx import Document

    _guard_size(data)
    try:
        doc = Document(io.BytesIO(data))
    except Exception as exc:
        raise UnreadableFile(
            "That .docx could not be opened — it may be corrupt, password-protected, "
            "or not really a Word file."
        ) from exc

    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(p for p in parts if p.strip())


def read_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover
        raise NeedsOCR("pypdf is not installed - run: pip install pypdf") from exc

    _guard_size(data)
    try:
        reader = PdfReader(io.BytesIO(data))
        pages = [(page.extract_text() or "") for page in reader.pages]
    except Exception as exc:
        raise UnreadableFile(
            "That PDF could not be read — it may be corrupt, encrypted, or not really a PDF."
        ) from exc

    text = "\n".join(pages).strip()
    if len(text) < 40:
        raise NeedsOCR("This PDF looks like a scan - no text layer found.")

    tables = pdf_tables(data)
    if tables:
        # A plain text extractor shreds a table's rows into word soup; the
        # reconstructed key-value lines below carry the SAME values with
        # their headers reattached. Clearly marked, so the user deletes the
        # section if the loose extraction already read fine.
        text += ("\n\n[TABLES RECONSTRUCTED - the same values as above with "
                 "their column headers reattached; delete this section if "
                 "the text above already reads correctly]\n"
                 + tables_to_text(tables))
    return text


def pdf_tables(data: bytes) -> list[list[list[str]]]:
    """
    Tables detected in the PDF, as rows of cells. Echo and lab reports keep
    their measurements (EF, LVIDd...) in grids that plain text extraction
    destroys. Empty list when pdfplumber is absent or finds nothing - the
    text path above never depends on this.
    """
    try:
        import pdfplumber
    except ImportError:
        return []
    tables: list[list[list[str]]] = []
    try:
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page in pdf.pages:
                for table in page.extract_tables() or []:
                    rows = [[(cell or "").strip() for cell in row]
                            for row in table if row]
                    rows = [r for r in rows if any(r)]
                    if len(rows) >= 2 and max(len(r) for r in rows) >= 2:
                        tables.append(rows)
    except Exception:
        return []  # a broken table layer must not break the text path
    return tables


def tables_to_text(tables: list[list[list[str]]]) -> str:
    """
    Rows of cells to readable lines.

    Two-column tables are key-value pairs ("EF: 62 %"). Wider tables get the
    header reattached to every cell ("LVIDd: 4.8 cm · LVIDs: 3.1 cm"), so a
    value never floats free of its meaning.
    """
    blocks: list[str] = []
    for table in tables:
        lines: list[str] = []
        width = max(len(row) for row in table)
        if width == 2:
            for row in table:
                key = (row[0] or "").strip()
                value = (row[1] if len(row) > 1 else "").strip()
                if key and value:
                    lines.append(f"{key}: {value}")
                elif key or value:
                    lines.append(key or value)
        else:
            header = [c.strip() for c in table[0]]
            has_header = sum(1 for c in header if c) >= 2 and not any(
                ch.isdigit() for c in header for ch in c
            )
            body = table[1:] if has_header else table
            for row in body:
                cells = []
                for i, cell in enumerate(row):
                    cell = (cell or "").strip()
                    if not cell:
                        continue
                    name = header[i] if has_header and i < len(header) and header[i] else ""
                    cells.append(f"{name}: {cell}" if name else cell)
                if cells:
                    lines.append(" · ".join(cells))
        if lines:
            blocks.append("\n".join(lines))
    return "\n\n".join(blocks)


def read_any(filename: str, data: bytes) -> str:
    """Dispatch on extension. Raises NeedsOCR when only vision can help."""
    lower = filename.lower()
    ext = lower[lower.rfind("."):] if "." in lower else ""

    _guard_size(data)
    if ext in TEXT_EXT:
        return read_text_file(data)
    if ext == ".docx":
        return read_docx(data)
    if ext == ".pdf":
        return read_pdf(data)
    if ext in IMAGE_EXT:
        raise NeedsOCR("Images have no text layer - AI OCR is required.")
    return read_text_file(data)


def mime_for(filename: str) -> str:
    lower = filename.lower()
    ext = lower[lower.rfind("."):] if "." in lower else ""
    return MIME_BY_EXT.get(ext, "application/octet-stream")
