"""
Offline self-test: parse the samples, build the .docx, run the word-loss audit
and assert every HC FORMAT rule on the actual XML that Word will read.

    python selftest.py
"""

from __future__ import annotations

import glob
import io
import os
import sys
import zipfile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

import templates
from hc_format import LINE_SPACING, Block, ParseOptions, Span, build_docx, parse_report
from verify import audit

HERE = os.path.dirname(os.path.abspath(__file__))
failures: list[str] = []


def check(condition: bool, message: str) -> None:
    if not condition:
        failures.append(message)


def inspect(docx_bytes: bytes, label: str) -> None:
    doc = Document(io.BytesIO(docx_bytes))

    for section in doc.sections:
        check(abs(section.left_margin.inches - 1.0) < 0.01, f"{label}: left margin not 1 inch")
        check(abs(section.top_margin.inches - 1.0) < 0.01, f"{label}: top margin not 1 inch")

    paragraphs = [p for p in doc.paragraphs if p.text.strip()]
    check(bool(paragraphs), f"{label}: document is empty")

    # Rule 2 - professional spacing: 1.5 line spacing on every report paragraph.
    for p in paragraphs:
        spacing = p.paragraph_format.line_spacing
        check(spacing == LINE_SPACING,
              f"{label}: line spacing is {spacing}, expected {LINE_SPACING} — {p.text[:35]!r}")

    # Rule 1 - Arial 12 pt black everywhere.
    for p in paragraphs:
        for run in p.runs:
            if not run.text.strip():
                continue
            check(run.font.name == "Arial", f"{label}: non-Arial run {run.text[:30]!r}")
            check(run.font.size is not None and run.font.size.pt == 12,
                  f"{label}: non-12pt run {run.text[:30]!r}")
            check(run.font.color.rgb is not None and str(run.font.color.rgb) == "000000",
                  f"{label}: non-black run {run.text[:30]!r}")

    # Rule 3 - title.
    title = paragraphs[0]
    check(title.alignment == WD_ALIGN_PARAGRAPH.CENTER, f"{label}: title not centred")
    check(all(r.font.bold for r in title.runs if r.text.strip()), f"{label}: title not bold")
    check(all(r.font.underline for r in title.runs if r.text.strip()), f"{label}: title not underlined")
    check(title.text == title.text.upper(), f"{label}: title not uppercase")

    # Rules 4/5/6 - headings, subheadings, impression bullets.
    saw_heading = saw_subheading = saw_bullet = saw_bold_bullet = False
    section_now = ""
    for p in paragraphs[1:]:
        runs = [r for r in p.runs if r.text.strip()]
        if not runs:
            continue
        bulleted = p.style.name == "List Bullet"
        bold = all(r.font.bold for r in runs)
        italic = all(r.font.italic for r in runs)
        underline = all(bool(r.font.underline) for r in runs)

        if bold and underline and not bulleted and p.text == p.text.upper():
            saw_heading = True
            section_now = p.text.rstrip(":").strip()
            check(p.alignment in (WD_ALIGN_PARAGRAPH.LEFT, None),
                  f"{label}: heading {p.text!r} not left aligned")
            continue

        if section_now in ("FINDINGS", "OBSERVATIONS"):
            if italic and underline:
                saw_subheading = True
                check(not bold, f"{label}: findings subheading {p.text!r} must NOT be bold")
            elif bulleted:
                saw_bullet = True
                check(not bold, f"{label}: findings bullet {p.text[:40]!r} should not be bold")

        if section_now in ("IMPRESSION", "CONCLUSION") and bulleted:
            saw_bold_bullet = True
            check(bold, f"{label}: impression bullet {p.text[:40]!r} must be BOLD")

    # Inline metadata headings: label run bold+underlined, value run plain.
    for p in paragraphs:
        runs = [r for r in p.runs if r.text.strip()]
        if len(runs) < 2 or not runs[0].font.bold:
            continue
        check(bool(runs[0].font.underline),
              f"{label}: inline heading {runs[0].text!r} not underlined")
        check(runs[0].text == runs[0].text.upper(),
              f"{label}: inline heading {runs[0].text!r} not uppercase")
        for tail in runs[1:]:
            check(not tail.font.bold,
                  f"{label}: value {tail.text[:30]!r} must not inherit heading bold")
            check(not tail.font.underline,
                  f"{label}: value {tail.text[:30]!r} must not inherit heading underline")

    check(saw_heading, f"{label}: no main heading detected")
    check(saw_subheading, f"{label}: no italic+underlined organ subheading detected")
    check(saw_bullet, f"{label}: no findings bullet detected")
    check(saw_bold_bullet, f"{label}: no bold impression bullet detected")


def check_as_is() -> None:
    """As-is mode must print the paste character for character."""
    raw = "IMPRESSION\n    indented line\n\n1) marker kept\nlowercase stays\n"
    blocks = parse_report(raw, ParseOptions(preserve_as_is=True)).blocks
    data = build_docx(blocks)
    doc = Document(io.BytesIO(data))
    got = [p.text for p in doc.paragraphs]

    check(got == ["IMPRESSION", "    indented line", "", "1) marker kept", "lowercase stays"],
          f"as-is: paragraphs differ from the source -> {got!r}")
    check(all(p.style.name != "List Bullet" for p in doc.paragraphs),
          "as-is: a bullet was added")
    check("1) marker kept" in got, "as-is: a list marker was stripped")
    check("lowercase stays" in got, "as-is: text was upper-cased")

    xml = zipfile.ZipFile(io.BytesIO(data)).read("word/document.xml").decode("utf8")
    check('xml:space="preserve">    indented' in xml,
          "as-is: leading indentation lost (missing xml:space=preserve)")

    report = audit(raw, data, preserve_as_is=True)
    check(report.ok, f"as-is: word-loss audit failed -> {report.summary}")


def check_rich_text() -> None:
    """Inline emphasis must add to the line's format, never replace it."""
    block = Block(
        kind="bold_bullet",
        text="Small CSF collection noted.",
        spans=[Span("Small CSF collection "), Span("noted.", italic=True, underline=True)],
    )
    doc = Document(io.BytesIO(build_docx([block])))
    runs = [r for r in doc.paragraphs[0].runs if r.text]
    check(len(runs) == 2, f"rich text: expected 2 runs, got {len(runs)}")
    check(all(r.font.bold for r in runs),
          "rich text: rule 6 bold was lost on an impression bullet")
    check(not runs[0].font.italic and bool(runs[1].font.italic),
          "rich text: italic did not land on the right run")
    check(bool(runs[1].font.underline), "rich text: underline did not land")
    check("".join(r.text for r in runs) == block.text,
          "rich text: run text does not reassemble to the block text")


def check_templates() -> None:
    """A doctor's template must drive font, size, spacing, alignment and emphasis."""
    tpl = templates.copy_of(templates.HC_FORMAT, "Dr. Test", doctor="Dr. Test")
    tpl.font_name = "Times New Roman"
    tpl.font_size = 11
    tpl.line_spacing = 1.0
    tpl.styles["title"].align = "left"
    tpl.styles["title"].underline = False
    tpl.styles["text"].align = "justify"

    blocks = [Block("title", "MRI BRAIN"), Block("text", "Body text.")]
    doc = Document(io.BytesIO(build_docx(blocks, template=tpl)))
    title, body = doc.paragraphs[0], doc.paragraphs[1]

    for p in (title, body):
        run = p.runs[0]
        check(run.font.name == "Times New Roman", f"template: font not applied to {p.text!r}")
        check(run.font.size.pt == 11, f"template: size not applied to {p.text!r}")
        check(p.paragraph_format.line_spacing == 1.0,
              f"template: line spacing not applied to {p.text!r}")

    check(title.alignment == WD_ALIGN_PARAGRAPH.LEFT, "template: title alignment not applied")
    check(not title.runs[0].font.underline, "template: title underline override ignored")
    check(body.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY, "template: body alignment not applied")

    # Round-trip through JSON, the way a saved doctor template is loaded.
    restored = templates.from_dict(templates.to_dict(tpl))
    check(restored.font_name == "Times New Roman" and restored.line_spacing == 1.0,
          "template: JSON round-trip lost the basics")
    check(restored.style("text").align == "justify",
          "template: JSON round-trip lost a per-line style")
    check(not restored.style("title").underline,
          "template: JSON round-trip lost a per-line override")

    # The built-in must stay exactly as signed off.
    hc = templates.HC_FORMAT
    check(hc.font_name == "Arial" and hc.font_size == 12 and hc.line_spacing == 1.5,
          "template: built-in HC FORMAT drifted")
    check(hc.style("bold_bullet").bold and hc.style("bold_bullet").bullet,
          "template: HC FORMAT impression bullets are no longer bold bullets")
    check(hc.style("subheading").italic and hc.style("subheading").underline
          and not hc.style("subheading").bold,
          "template: HC FORMAT organ subheadings drifted")


def check_drafting() -> None:
    """Few-shot prompt assembly and the measurement guard. No network."""
    import ai_parser

    tpl = templates.copy_of(templates.HC_FORMAT, "Dr. Voice", doctor="Dr. A Voice")
    tpl.style_notes = "Always number the impression."
    tpl.examples = [
        "USG ABDOMEN\nIMPRESSION:\n1. Cholelithiasis.\nPlease correlate clinically.",
        "USG PELVIS\nIMPRESSION:\n1. Normal study.\nPlease correlate clinically.",
    ]

    prompt = ai_parser.build_draft_prompt(tpl, "mult gallstones 4-11mm", "the IMPRESSION only")
    check("Dr. A Voice" in prompt, "draft prompt: doctor name missing")
    check("Always number the impression." in prompt, "draft prompt: style notes missing")
    check("EXAMPLE 1" in prompt and "EXAMPLE 2" in prompt,
          "draft prompt: past reports not used as few-shot examples")
    check("Cholelithiasis." in prompt, "draft prompt: example body missing")
    check("mult gallstones 4-11mm" in prompt, "draft prompt: the notes themselves are missing")
    check("the IMPRESSION only" in prompt, "draft prompt: section not passed through")

    empty = templates.copy_of(templates.HC_FORMAT, "Dr. New", doctor="Dr. New")
    bare = ai_parser.build_draft_prompt(empty, "some notes")
    check("Nothing is on file" in bare,
          "draft prompt: no fallback when a doctor has nothing on file")
    check("EXAMPLE" not in bare, "draft prompt: invented an example out of nothing")

    # The guard exists to catch exactly one failure: a lost or altered measurement.
    check(ai_parser.missing_facts("liver 14.2 cm, 3 stones", "Liver measures 14.2 cm. 3 stones.") == [],
          "measurement guard: false alarm on an intact draft")
    check(ai_parser.missing_facts("liver 14.2 cm", "Liver measures 15.6 cm.") == ["14.2"],
          "measurement guard: missed a changed measurement")
    check(ai_parser.missing_facts("4 mm to 11 mm", "4 mm lesion") == ["11"],
          "measurement guard: missed a dropped measurement")


def check_learning() -> None:
    """Corrections must turn into rules, persist, and reach the next prompt."""
    import ai_parser

    tpl = templates.copy_of(templates.HC_FORMAT, "Dr. Learn", doctor="Dr. Learn")

    # A correction plus the rules distilled from it.
    tpl = templates.remember_correction(
        tpl,
        before="IMPRESSION:\n- Multiple stones in the gallbladder.",
        after="IMPRESSION:\n1. Multiple calculi in the gallbladder.",
        note="calculi, and number the impression",
        rules=["Write 'calculi' rather than 'stones'.", "Number the impression points."],
    )
    check(len(tpl.corrections) == 1, "learning: the correction was not recorded")
    check(len(tpl.preferences) == 2, "learning: the rules were not recorded")
    check(tpl.corrections[0].when, "learning: the correction has no timestamp")

    # The same lesson learned twice must not weight the prompt twice.
    tpl = templates.remember_correction(
        tpl, before="x", after="y", rules=["write 'calculi' rather than 'stones'"]
    )
    check(len(tpl.preferences) == 2, f"learning: duplicate rule stored ({tpl.preferences})")

    # An edit that changes nothing is not a correction.
    same = templates.remember_correction(tpl, before="same", after="same")
    check(len(same.corrections) == len(tpl.corrections),
          "learning: a no-op edit was stored as a correction")

    # Answers are remembered so the doctor is not asked twice.
    tpl = templates.remember_answer(tpl, "Which side is the kidney?", "Right kidney")
    check(tpl.answered.get("Which side is the kidney?") == "Right kidney",
          "learning: the answer was not remembered")

    # Everything must survive a save/load round-trip.
    restored = templates.from_dict(templates.to_dict(tpl))
    check(len(restored.corrections) == len(tpl.corrections),
          "learning: corrections lost in the JSON round-trip")
    check(restored.preferences == tpl.preferences,
          "learning: rules lost in the JSON round-trip")
    check(restored.answered == tpl.answered,
          "learning: answers lost in the JSON round-trip")
    check(restored.corrections[0].note == "calculi, and number the impression",
          "learning: the reason given was lost in the round-trip")

    # …and must actually reach the model.
    prompt = ai_parser.build_draft_prompt(restored, "gallstones seen")
    check("Rules learned from" in prompt, "learning: rules never reach the prompt")
    check("Write 'calculi' rather than 'stones'." in prompt,
          "learning: a specific rule never reaches the prompt")
    check("CORRECTION 1" in prompt, "learning: before/after pairs never reach the prompt")
    check("Multiple calculi in the gallbladder." in prompt,
          "learning: the corrected text never reaches the prompt")
    check("do not ask these again" in prompt,
          "learning: answered questions never reach the prompt")
    check("Right kidney" in prompt, "learning: a saved answer never reaches the prompt")

    # Caps keep the prompt from growing without limit.
    big = templates.copy_of(templates.HC_FORMAT, "Dr. Many")
    for i in range(templates.MAX_EXAMPLES + 5):
        big = templates.remember_example(big, f"report number {i}")
    check(len(big.examples) == templates.MAX_EXAMPLES,
          f"learning: examples not capped ({len(big.examples)})")
    check("report number 12" in big.examples, "learning: the cap dropped the newest example")

    for i in range(templates.MAX_CORRECTIONS + 5):
        big = templates.remember_correction(big, before=f"b{i}", after=f"a{i}")
    check(len(big.corrections) == templates.MAX_CORRECTIONS,
          f"learning: corrections not capped ({len(big.corrections)})")

    # Forgetting must work - a wrong lesson has to be removable.
    pruned = templates.forget_preference(restored, "Number the impression points.")
    check("Number the impression points." not in pruned.preferences,
          "learning: a rule could not be forgotten")
    check(len(pruned.preferences) == len(restored.preferences) - 1,
          "learning: forgetting removed the wrong number of rules")


def check_dictation() -> None:
    """Misheard words must be remembered and must reach the next transcription."""
    import ai_parser

    tpl = templates.copy_of(templates.HC_FORMAT, "Dr. Voice", doctor="Dr. Voice")

    # Nothing on file: the prompt must say so rather than pretend.
    bare = ai_parser.build_transcribe_prompt(tpl)
    check("Nothing is on file" in bare, "dictation: no fallback for an unknown voice")
    check("Dr. Voice" in bare, "dictation: the doctor is not named in the prompt")

    # A mishearing is recorded, and the corrected word becomes vocabulary.
    tpl = templates.remember_dictation_fix(tpl, "colic list", "cholelithiasis")
    check(len(tpl.dictation_fixes) == 1, "dictation: the mishearing was not recorded")
    check("cholelithiasis" in tpl.vocabulary,
          "dictation: the corrected word did not become vocabulary")
    check(tpl.dictation_fixes[0].when, "dictation: the fix has no timestamp")

    # A no-op fix is not a fix.
    same = templates.remember_dictation_fix(tpl, "liver", "liver")
    check(len(same.dictation_fixes) == 1, "dictation: a no-op fix was stored")
    empty = templates.remember_dictation_fix(tpl, "", "something")
    check(len(empty.dictation_fixes) == 1, "dictation: an empty fix was stored")

    # The same mishearing twice must not be stored twice.
    twice = templates.remember_dictation_fix(tpl, "Colic List", "Cholelithiasis")
    check(len(twice.dictation_fixes) == 1,
          f"dictation: duplicate fix stored ({len(twice.dictation_fixes)})")

    tpl = templates.remember_vocabulary(tpl, ["hydronephrosis", "hydronephrosis", "craniocaudal"])
    check(tpl.vocabulary.count("hydronephrosis") == 1, "dictation: duplicate term stored")
    check("craniocaudal" in tpl.vocabulary, "dictation: a new term was dropped")

    # Everything survives the round-trip to disk…
    restored = templates.from_dict(templates.to_dict(tpl))
    check(restored.vocabulary == tpl.vocabulary,
          "dictation: vocabulary lost in the JSON round-trip")
    check(len(restored.dictation_fixes) == len(tpl.dictation_fixes),
          "dictation: fixes lost in the JSON round-trip")
    check(restored.dictation_fixes[0].after == "cholelithiasis",
          "dictation: the corrected word was mangled in the round-trip")

    # …and reaches the next transcription.
    prompt = ai_parser.build_transcribe_prompt(restored, context="USG abdomen")
    check("cholelithiasis" in prompt, "dictation: learned vocabulary never reaches the prompt")
    check('heard "colic list"' in prompt, "dictation: past mishearings never reach the prompt")
    check("USG abdomen" in prompt, "dictation: the study context never reaches the prompt")
    check("Nothing is on file" not in prompt,
          "dictation: still claims nothing is on file after learning")

    # Caps.
    many = templates.remember_vocabulary(
        tpl, [f"term{i}" for i in range(templates.MAX_VOCABULARY + 20)]
    )
    check(len(many.vocabulary) == templates.MAX_VOCABULARY,
          f"dictation: vocabulary not capped ({len(many.vocabulary)})")

    pruned = templates.forget_vocabulary(restored, "cholelithiasis")
    check("cholelithiasis" not in pruned.vocabulary, "dictation: a term could not be forgotten")


def check_speech_backends() -> None:
    """The AI4Bharat backend: config surface and failure messages. No network."""
    import ai_parser
    import speech

    check(set(speech.ENGINES) == {"gemini", "ai4bharat+gemini", "ai4bharat"},
          f"speech: unexpected engine list {list(speech.ENGINES)}")
    check("en" in speech.LANGUAGES and "hi" in speech.LANGUAGES,
          "speech: English and Hindi must both be offered")
    check(len(speech.LANGUAGES) >= 23,
          f"speech: expected the 22 scheduled languages plus English, got {len(speech.LANGUAGES)}")
    check(all(m.startswith("ai4bharat/") for m in speech.AI4BHARAT_PRESETS),
          "speech: a preset is not an AI4Bharat repo")

    # A missing token must be explained, not raised as a stray urllib error.
    try:
        speech.transcribe_ai4bharat_hf(b"x", "audio/wav", model="ai4bharat/indicwhisper",
                                       hf_token="")
        check(False, "speech: a missing HF token was not caught")
    except speech.SpeechError as exc:
        check("HF_TOKEN" in str(exc), f"speech: unhelpful message for a missing token: {exc}")

    # WAV passes through untouched; nothing is silently mangled.
    same, mime = speech.to_wav(b"RIFFfake", "audio/wav")
    check(same == b"RIFFfake" and mime == "audio/wav",
          "speech: WAV input was altered on the way through")

    # The layout pass must carry the doctor's vocabulary, like the audio path does.
    tpl = templates.copy_of(templates.HC_FORMAT, "Dr. Indic", doctor="Dr. Indic")
    tpl = templates.remember_dictation_fix(tpl, "colic list", "cholelithiasis")
    prompt = ai_parser.build_transcribe_prompt(tpl, context="USG abdomen")
    check("cholelithiasis" in prompt,
          "speech: vocabulary must reach the layout pass as well as the audio pass")
    # Wrapping in the prompt text moves, so assert on the words, not the line breaks.
    layout_prompt = " ".join(ai_parser.STRUCTURE_DICTATION_PROMPT.split())
    check("Do not translate" in layout_prompt,
          "speech: the layout pass must be told not to translate an Indic-language report")
    check("keep the report in THAT language" in layout_prompt,
          "speech: the layout pass must be told to keep the doctor's language")
    check("NEVER add a finding" in layout_prompt and "NEVER invent a digit" in layout_prompt,
          "speech: the layout pass lost its safety rules")


def main() -> int:
    paths = sorted(glob.glob(os.path.join(HERE, "samples", "*.txt")))
    if not paths:
        print("No samples found in samples/")
        return 1

    for path in paths:
        label = os.path.basename(path)
        with open(path, encoding="utf-8") as fh:
            raw = fh.read()

        result = parse_report(raw, ParseOptions())
        docx_bytes = build_docx(result.blocks)
        inspect(docx_bytes, label)

        report = audit(raw, docx_bytes)
        check(report.ok, f"{label}: word-loss audit failed -> {report.summary}")

        out = os.path.join(HERE, "output", label.replace(".txt", "_HC_Format.docx"))
        os.makedirs(os.path.dirname(out), exist_ok=True)
        with open(out, "wb") as fh:
            fh.write(docx_bytes)

        status = "PASS" if report.ok else "FAIL"
        print(f"{label}: {len(result.blocks)} blocks, audit {status} "
              f"({report.source_tokens} tokens) -> {os.path.relpath(out, HERE)}")
        if not report.ok:
            print(f"   missing: {report.missing[:12]}")
            print(f"   added:   {report.added[:12]}")

    check_as_is()
    check_rich_text()
    check_templates()
    check_drafting()
    check_learning()
    check_dictation()
    check_speech_backends()
    print("as-is, rich text, templates, drafting, learning, dictation and AI4Bharat: checked")

    print()
    if failures:
        print(f"{len(failures)} check(s) FAILED:")
        for message in failures:
            print("  -", message)
        return 1
    print("All HC FORMAT checks passed.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
