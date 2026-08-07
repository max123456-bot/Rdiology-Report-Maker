"""
Whole-environment check: does every part still work when wired to every other part?

selftest.py tests each feature alone. security_check.py attacks each one. This
walks the journeys a clinic actually takes, end to end, and asserts the handoffs
between features - the seams where things quietly break.

    python integration_check.py

Offline. No API key, no network.
"""

from __future__ import annotations

import io
import os
import sys
import zipfile

from docx import Document


import readers
import storage
import templates
import verify
from hc_format import Block, ParseOptions, Span, build_docx, parse_report

failures: list[str] = []
PROBE = "__integration_doctor"


def check(condition: bool, message: str) -> None:
    print(("  ok    " if condition else "  FAIL  ") + message)
    if not condition:
        failures.append(message)


REPORT = """ULTRASONOGRAPHY OF WHOLE ABDOMEN AND PELVIS

PATIENT NAME: Mrs. Sunita Devi
AGE/SEX: 41 Y / F
CLINICAL HISTORY: Right upper quadrant pain since 2 weeks.
TECHNIQUE: Real time grey scale sonography with a 3.5 MHz curvilinear transducer.

FINDINGS:
Liver: Normal in size measuring 14.2 cm in craniocaudal span. No focal lesion is seen.
Gallbladder: Multiple mobile echogenic foci measuring 4 mm to 11 mm within the lumen.
Kidneys: Right kidney measures 9.8 x 4.4 cm. No calculus or hydronephrosis is seen.

IMPRESSION:
1. Multiple cholelithiasis measuring 4 mm to 11 mm.
2. No hepatic or renal abnormality detected.

Please correlate clinically.
DR. ABHIJEET SHARMA
"""


def journey_default_format() -> None:
    print("\n1. Boss sends a report -> HC FORMAT .docx, audit passes")
    result = parse_report(REPORT, ParseOptions())
    docx = build_docx(result.blocks)
    report = verify.audit(REPORT, docx)
    check(report.ok, f"word-loss audit passes ({report.source_tokens} tokens)")

    doc = Document(io.BytesIO(docx))
    paragraphs = [p for p in doc.paragraphs if p.text.strip()]
    check(paragraphs[0].text == paragraphs[0].text.upper(), "title upper-cased")
    check(all(p.paragraph_format.line_spacing == 1.5 for p in paragraphs),
          "1.5 line spacing on every paragraph")
    check(any(p.style.name == "List Bullet" for p in paragraphs), "bullets present")


def journey_doctor_template() -> None:
    print("\n2. New doctor template drives the same report")
    tpl = templates.copy_of(templates.HC_FORMAT, PROBE, doctor="Dr. Integration")
    tpl.font_name = "Times New Roman"
    tpl.font_size = 11
    tpl.line_spacing = 1.15
    tpl.letterhead_name = "Integration Imaging"
    templates.save(tpl)

    loaded = templates.load_all().get(PROBE)
    check(loaded is not None, "template round-trips through disk")

    blocks = parse_report(REPORT, ParseOptions()).blocks
    docx = build_docx(blocks, template=loaded,
                      letterhead={"name": loaded.letterhead_name}, page_numbers=True)
    doc = Document(io.BytesIO(docx))
    body = [p for p in doc.paragraphs if p.text.strip()]
    fonts = {r.font.name for p in body for r in p.runs if r.text.strip()}
    check(fonts == {"Times New Roman"}, f"every run uses the doctor's font (saw {fonts})")

    # The letterhead is deliberately single-spaced header chrome; the report body
    # below it must carry the doctor's spacing.
    title_at = next(i for i, p in enumerate(body) if "ULTRASONOGRAPHY" in p.text)
    header, report_body = body[:title_at], body[title_at:]
    check(all(p.paragraph_format.line_spacing == 1.15 for p in report_body),
          "the doctor's line spacing is applied to the report body")
    check(all(p.paragraph_format.line_spacing == 1.0 for p in header),
          "the letterhead stays tight regardless of the report spacing")

    # The letterhead and the page-number footer must not read as word loss.
    report = verify.audit(REPORT, docx, letterhead_text=loaded.letterhead_name,
                          page_numbers=True)
    check(report.ok, f"audit still passes with letterhead + page numbers ({report.summary})")


def journey_as_is_with_template() -> None:
    print("\n3. As-is mode still obeys the doctor's template")
    tpl = templates.load_all()[PROBE]
    raw = "IMPRESSION\n    indented line\n\n1) marker kept\nlowercase stays"
    blocks = parse_report(raw, ParseOptions(preserve_as_is=True)).blocks
    docx = build_docx(blocks, template=tpl)
    doc = Document(io.BytesIO(docx))
    got = [p.text for p in doc.paragraphs]
    check(got == ["IMPRESSION", "    indented line", "", "1) marker kept", "lowercase stays"],
          "text printed exactly as pasted")
    fonts = {r.font.name for p in doc.paragraphs for r in p.runs if r.text.strip()}
    check(fonts == {"Times New Roman"}, "as-is text still uses the doctor's font")
    check(verify.audit(raw, docx, preserve_as_is=True).ok, "as-is audit passes")


def journey_edit_then_learn() -> None:
    print("\n4. Doctor edits the preview, then the app learns from it")
    blocks = parse_report(REPORT, ParseOptions()).blocks

    # Simulate the WYSIWYG: change a measurement and italicise a word.
    edited = []
    for b in blocks:
        if "14.2" in b.text:
            edited.append(Block(kind=b.kind, text=b.text.replace("14.2", "15.6"),
                                trailer=b.trailer))
        elif b.kind == "bold_bullet" and "cholelithiasis" in b.text:
            head, _, tail = b.text.partition("cholelithiasis")
            edited.append(Block(kind=b.kind, text=b.text, trailer=b.trailer,
                                spans=[Span(head), Span("cholelithiasis", italic=True),
                                       Span(tail)]))
        else:
            edited.append(b)

    docx = build_docx(edited)
    report = verify.audit(REPORT, docx)
    check(not report.ok, "the audit notices the doctor's edit rather than passing silently")
    check(any(tok == "15.6" for tok, _ in report.added), "the new measurement is reported as added")
    check(any(tok == "14.2" for tok, _ in report.missing), "the old measurement is reported as missing")

    doc = Document(io.BytesIO(docx))
    target = [p for p in doc.paragraphs if "cholelithiasis" in p.text and p.style.name == "List Bullet"]
    check(bool(target), "the impression bullet survived the span edit")
    if target:
        runs = [r for r in target[0].runs if r.text.strip()]
        check(all(r.font.bold for r in runs), "rule 6 bold survived inline italics")
        check(any(r.font.italic for r in runs), "the inline italic landed")

    # And the learning loop stores it.
    tpl = templates.load_all()[PROBE]
    taught = templates.remember_correction(
        tpl, "Multiple stones.", "1. Multiple calculi.",
        rules=["Write 'calculi' rather than 'stones'."])
    taught = templates.remember_dictation_fix(taught, "colic list", "cholelithiasis")
    templates.save(taught)
    back = templates.load_all()[PROBE]
    check(back.preferences and back.vocabulary and back.corrections,
          "corrections, rules and vocabulary all persist together")


def journey_prompts_carry_everything() -> None:
    print("\n5. Everything learned reaches every prompt that should see it")
    import ai_parser

    tpl = templates.load_all()[PROBE]
    draft = ai_parser.build_draft_prompt(tpl, "gallstones seen")
    check("Write 'calculi' rather than 'stones'." in draft, "learned rules reach the drafting prompt")
    check("CORRECTION 1" in draft, "before/after pairs reach the drafting prompt")

    listen = ai_parser.build_transcribe_prompt(tpl, context="USG abdomen")
    check("cholelithiasis" in listen, "vocabulary reaches the transcription prompt")
    check('heard "colic list"' in listen, "past mishearings reach the transcription prompt")
    check("USG abdomen" in listen, "the study context reaches the transcription prompt")


def journey_batch() -> None:
    print("\n6. Batch: many reports -> one ZIP, each audited")
    tpl = templates.load_all()[PROBE]
    jobs = [("a", REPORT), ("b", REPORT.replace("Sunita Devi", "Anita Rao")),
            ("c", "X-RAY CHEST\n\nFINDINGS:\nLung fields are clear.\n\nIMPRESSION:\nNormal study.")]
    buf = io.BytesIO()
    rows = []
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for i, (label, text) in enumerate(jobs, start=1):
            blocks = parse_report(text, ParseOptions()).blocks
            docx = build_docx(blocks, template=tpl)
            name = f"report_{i}.docx"
            zf.writestr(name, docx)
            rows.append(verify.audit(text, docx).ok)

    check(all(rows), "every report in the batch passed its audit")
    with zipfile.ZipFile(io.BytesIO(buf.getvalue())) as zf:
        names = zf.namelist()
        check(len(names) == len(jobs), f"{len(names)} files in the ZIP")
        check(len(set(names)) == len(names), "no two reports overwrote each other")
        first = Document(io.BytesIO(zf.read(names[0])))
        check(any(r.font.name == "Times New Roman"
                  for p in first.paragraphs for r in p.runs if r.text.strip()),
              "batch output uses the selected doctor's template")


def journey_uploads() -> None:
    print("\n7. Upload paths feed the same pipeline")
    docx_in = build_docx(parse_report(REPORT, ParseOptions()).blocks)
    text = readers.read_any("report.docx", docx_in)
    check("cholelithiasis" in text.lower(), "a .docx upload reads back its text")

    blocks = parse_report(text, ParseOptions()).blocks
    rebuilt = build_docx(blocks)
    check(verify.audit(text, rebuilt).ok, "re-formatting an uploaded .docx passes the audit")

    check(readers.read_any("x.txt", REPORT.encode("utf-8")).startswith("ULTRASONOGRAPHY"),
          "a .txt upload reads back")


def journey_template_deletion() -> None:
    print("\n8. Deleting a doctor leaves the rest of the app standing")
    check(templates.delete(PROBE), "the probe template was deleted")
    remaining = templates.load_all()
    check(PROBE not in remaining, "it is gone from the list")
    check(templates.HC_FORMAT.name in remaining, "the built-in survives")
    docx = build_docx(parse_report(REPORT, ParseOptions()).blocks,
                      template=remaining[templates.HC_FORMAT.name])
    check(verify.audit(REPORT, docx).ok, "formatting still works after the delete")


def journey_save_safety() -> None:
    print("\n8b. Saving: backups, and a concurrent edit is refused not clobbered")
    name = "__integration_save"
    try:
        first = templates.copy_of(templates.HC_FORMAT, name, doctor="First")
        templates.save(first)
        stamp = templates.fingerprint(name)

        second = templates.copy_of(templates.load_all()[name], name, doctor="Second")
        templates.save(second, expect=stamp)
        check(templates.load_all()[name].doctor == "Second", "a matching save goes through")

        third = templates.copy_of(templates.load_all()[name], name, doctor="Third")
        try:
            templates.save(third, expect=stamp)  # stale: someone else already wrote
            check(False, "a stale write was allowed to clobber the newer version")
        except templates.ConflictError:
            check(True, "a stale write is refused rather than clobbering")
        check(templates.load_all()[name].doctor == "Second",
              "the refused write left the file untouched")

        backups = []
        if os.path.isdir(storage.get_store().backups):
            stem = os.path.basename(storage.get_store()._path(name))[:-5]
            backups = [f for f in os.listdir(storage.get_store().backups) if f.startswith(stem + ".")]
        check(bool(backups), f"a backup was kept before overwriting ({len(backups)} on file)")
    finally:
        templates.delete(name)
        if os.path.isdir(storage.get_store().backups):
            stem = os.path.basename(storage.get_store()._path(name))[:-5]
            for f in os.listdir(storage.get_store().backups):
                if f.startswith(stem + "."):
                    os.remove(os.path.join(storage.get_store().backups, f))


def journey_builtin_untouched() -> None:
    print("\n9. The signed-off HC FORMAT has not drifted")
    hc = templates.load_all()[templates.HC_FORMAT.name]
    check(hc.font_name == "Arial" and hc.font_size == 12, "Arial 12 pt")
    check(hc.line_spacing == 1.5, "1.5 line spacing")
    check(all(hc.margin_top == m for m in (hc.margin_bottom, hc.margin_left, hc.margin_right))
          and hc.margin_top == 1.0, "1 inch margins all round")
    check(hc.style("title").align == "center" and hc.style("title").bold
          and hc.style("title").underline and hc.style("title").uppercase, "rule 3 intact")
    check(hc.style("heading").bold and hc.style("heading").underline
          and hc.style("heading").uppercase, "rule 4 intact")
    check(hc.style("subheading").italic and hc.style("subheading").underline
          and not hc.style("subheading").bold, "rule 5 intact")
    check(hc.style("bold_bullet").bold and hc.style("bold_bullet").bullet, "rule 6 intact")
    check(hc.builtin, "still marked read-only")


def main() -> int:
    print("Integration — the journeys a clinic actually takes")
    try:
        journey_default_format()
        journey_doctor_template()
        journey_as_is_with_template()
        journey_edit_then_learn()
        journey_prompts_carry_everything()
        journey_batch()
        journey_uploads()
        journey_save_safety()
        journey_template_deletion()
        journey_builtin_untouched()
    finally:
        templates.delete(PROBE)

    print("\n" + "=" * 70)
    if failures:
        print(f"{len(failures)} integration failure(s):")
        for message in failures:
            print("  -", message)
        return 1
    print("Every journey completed and every handoff held.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
