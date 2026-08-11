"""
Word-loss audit.

Reads the generated .docx back off disk (not the in-memory Block list, so a
rendering bug cannot hide from the check) and compares its tokens against the
raw input the user pasted.  Case is ignored because HC FORMAT mandates
UPPERCASE titles and headings; bullet glyphs are ignored because Word supplies
its own.  Everything else - every word, every measurement, every number - must
match exactly, in count.
"""

from __future__ import annotations

import io
import re
from collections import Counter
from dataclasses import dataclass, field

from docx import Document

from hc_format import _BULLET_PREFIX

# Words / numbers / measurements. Keeps decimals and hyphenated terms together.
_TOKEN = re.compile(r"[A-Za-z]+(?:['\-][A-Za-z]+)*|\d+(?:[.,]\d+)*")
_NUMERIC = re.compile(r"^\d")
# Glyphs a list marker may have used in the source; Word renders bullets itself.
_IGNORED = {"•", "●", "▪", "‣", "⁃", "·", "*", "-", "–", "—"}


@dataclass
class AuditResult:
    ok: bool = False
    source_tokens: int = 0
    output_tokens: int = 0
    missing: list[tuple[str, int]] = field(default_factory=list)  # in source, short in output
    added: list[tuple[str, int]] = field(default_factory=list)  # in output, absent from source
    numbers_ok: bool = True
    missing_numbers: list[str] = field(default_factory=list)

    @property
    def summary(self) -> str:
        if self.ok:
            return (
                f"PASS - all {self.source_tokens} words and numbers preserved verbatim "
                f"(case-insensitive; bullet glyphs excluded)."
            )
        bits = []
        if self.missing:
            bits.append(f"{sum(c for _, c in self.missing)} token(s) missing from output")
        if self.added:
            bits.append(f"{sum(c for _, c in self.added)} token(s) added")
        if not self.numbers_ok:
            bits.append(f"{len(self.missing_numbers)} measurement/number mismatch")
        return "FAIL - " + "; ".join(bits)


def _tokenize(text: str) -> Counter:
    return Counter(t.lower() for t in _TOKEN.findall(text) if t not in _IGNORED)


def _strip_list_markers(text: str) -> str:
    """
    Drop the same leading list markers the parser drops ("1.", "a)", "-", "•").

    Word draws its own bullet, so a source "1." is presentation, not content.
    The regex demands whitespace after the marker, so a measurement that opens a
    line ("2.4 x 1.9 cm lesion...") is left untouched.
    """
    return "\n".join(_BULLET_PREFIX.sub("", line) for line in text.split("\n"))


def docx_text(docx_bytes: bytes) -> str:
    """Flatten a .docx back to plain text, including headers and tables."""
    doc = Document(io.BytesIO(docx_bytes))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    for section in doc.sections:
        parts.extend(p.text for p in section.header.paragraphs)
        parts.extend(p.text for p in section.footer.paragraphs)
    return "\n".join(parts)


def audit(
    raw_text: str,
    docx_bytes: bytes,
    *,
    letterhead_text: str = "",
    page_numbers: bool = False,
    preserve_as_is: bool = False,
) -> AuditResult:
    """
    Compare raw input against the rendered document. No word may go missing.

    In as-is mode nothing is stripped on either side, because the engine did not
    strip anything either - a "1." that opened a line is still a "1." in the
    output, so it must still be counted.
    """
    src = _tokenize(raw_text if preserve_as_is else _strip_list_markers(raw_text))
    out = _tokenize(docx_text(docx_bytes))

    # Text this app injected on purpose (letterhead, sign-off, "Page x of y").
    injected = letterhead_text + (" Page of" if page_numbers else "")
    for token, count in _tokenize(injected).items():
        out[token] = max(0, out[token] - count)

    missing = [(tok, src[tok] - out.get(tok, 0)) for tok in src if src[tok] > out.get(tok, 0)]
    added = [(tok, out[tok] - src.get(tok, 0)) for tok in out if out[tok] > src.get(tok, 0)]

    missing_numbers = sorted(
        {tok for tok, _ in missing if _NUMERIC.match(tok)}
    )

    result = AuditResult(
        source_tokens=sum(src.values()),
        output_tokens=sum(out.values()),
        missing=sorted(missing, key=lambda x: -x[1]),
        added=sorted(added, key=lambda x: -x[1]),
        numbers_ok=not missing_numbers,
        missing_numbers=missing_numbers,
    )
    result.ok = not result.missing and not result.added
    return result


# --------------------------------------------------------------------------- #
# Reconciliation - where exactly did the lost words live?
# --------------------------------------------------------------------------- #


@dataclass
class Insertion:
    """One dropped span, with enough context to put it back."""

    tokens: list[str]        # the words, verbatim from the source
    word_index: int          # position in the source's word sequence
    before: str              # up to three source words preceding the span
    after: str               # up to three source words following the span

    @property
    def text(self) -> str:
        return " ".join(self.tokens)


def reconciliation_plan(raw_text: str, output_text: str) -> list[Insertion]:
    """
    Word-level alignment between the source and any derived text, returning
    each dropped span with its exact position and context.

    difflib's SequenceMatcher gives the same longest-match alignment
    diff-match-patch would at word granularity, with nothing to install.
    Comparison is case-insensitive (headings get uppercased by design) and
    list markers are ignored on both sides.
    """
    source_words = _TOKEN.findall(_strip_list_markers(raw_text))
    output_words = _TOKEN.findall(_strip_list_markers(output_text))
    import difflib

    matcher = difflib.SequenceMatcher(
        None,
        [w.lower() for w in source_words],
        [w.lower() for w in output_words],
        autojunk=False,
    )
    plan: list[Insertion] = []
    for op, s1, s2, _t1, _t2 in matcher.get_opcodes():
        if op not in ("delete", "replace"):
            continue
        dropped = [w for w in source_words[s1:s2] if w not in _IGNORED]
        if not dropped:
            continue
        plan.append(Insertion(
            tokens=dropped,
            word_index=s1,
            before=" ".join(source_words[max(0, s1 - 3):s1]),
            after=" ".join(source_words[s2:s2 + 3]),
        ))
    return plan


def auto_reconcile(target_text: str, plan: list[Insertion]) -> tuple[str, list[Insertion], list[Insertion]]:
    """
    Re-insert dropped spans into `target_text` at their context anchors.

    Non-destructive and non-inventive: every inserted word comes verbatim
    from the plan (i.e. from the original source), and anything whose anchor
    cannot be found is returned as skipped rather than guessed at. Callers
    surface this behind an explicit user action - the engine never rewrites
    a report on its own; that is this project's first rule.
    """
    applied: list[Insertion] = []
    skipped: list[Insertion] = []
    text = target_text

    for insertion in plan:
        anchor = insertion.before.strip()
        if anchor:
            pattern = re.compile(
                r"(" + r"\s+".join(re.escape(w) for w in anchor.split()) + r")",
                re.I,
            )
            match = pattern.search(text)
            if match:
                text = (text[:match.end()] + " " + insertion.text
                        + text[match.end():])
                applied.append(insertion)
                continue
        after_anchor = insertion.after.strip()
        if after_anchor:
            pattern = re.compile(
                r"(" + r"\s+".join(re.escape(w) for w in after_anchor.split()) + r")",
                re.I,
            )
            match = pattern.search(text)
            if match:
                text = (text[:match.start()] + insertion.text + " "
                        + text[match.start():])
                applied.append(insertion)
                continue
        skipped.append(insertion)
    return text, applied, skipped


# --------------------------------------------------------------------------- #
# Attestation - a hash-chained record that the audit ran and what it said
# --------------------------------------------------------------------------- #


def _attest_key() -> bytes:
    """ATTEST_KEY from secrets/environment, or b'' when unset."""
    import os

    value = ""
    try:
        import streamlit as st

        if "ATTEST_KEY" in st.secrets:
            value = str(st.secrets["ATTEST_KEY"]).strip()
    except Exception:
        pass
    value = value or os.environ.get("ATTEST_KEY", "").strip()
    return value.encode("utf-8")


def hmac_signature(text: str) -> str:
    """
    HMAC-SHA256 over the text with ATTEST_KEY - non-repudiation for a signed
    report: only a holder of the key can produce it, so a matching signature
    proves both that the text is untampered AND that this system signed it.
    Returns "" when no key is configured - honest absence, never a fake.
    """
    import hashlib
    import hmac as _hmac

    key = _attest_key()
    if not key:
        return ""
    return _hmac.new(key, (text or "").encode("utf-8"), hashlib.sha256).hexdigest()


def verify_signature(text: str, signature: str) -> bool:
    """Constant-time check of a report's keyed signature."""
    import hmac as _hmac

    expected = hmac_signature(text)
    if not expected or not signature:
        return False
    return _hmac.compare_digest(expected, signature)


def attestation(raw_text: str, docx_bytes: bytes, ok: bool,
                previous_chain: str = "") -> dict:
    """
    A tamper-evident record of one audit: SHA-256 of the source, SHA-256 of
    the output, the verdict, and a chain hash binding this record to the one
    before it. Recompute the chain over the log and any edited entry breaks
    every hash after it - that is the compliance property.
    """
    import hashlib
    from datetime import datetime, timezone

    source_sha = hashlib.sha256((raw_text or "").encode("utf-8")).hexdigest()
    output_sha = hashlib.sha256(docx_bytes or b"").hexdigest()
    verdict = "PASS" if ok else "FAIL"
    chain = hashlib.sha256(
        f"{previous_chain}|{source_sha}|{output_sha}|{verdict}".encode("ascii")
    ).hexdigest()
    record = {
        "when": datetime.now(timezone.utc).isoformat(timespec="seconds"),
        "source_sha256": source_sha,
        "output_sha256": output_sha,
        "verdict": verdict,
        "chain": chain,
    }
    # With ATTEST_KEY configured the chain itself is HMAC-signed, upgrading
    # tamper-EVIDENT (anyone can recompute a plain hash chain over altered
    # data) to tamper-PROOF against anyone who does not hold the key.
    signed = hmac_signature(chain)
    if signed:
        record["signature"] = signed
    return record


def audit_chain_status(entries: list[dict]) -> dict:
    """
    Recompute the attestation hash chain over audit-log entries, oldest first.

    Each link is checked against its own recorded predecessor rather than the
    recomputed one, so a single tampered record reads as ONE broken link, not
    as everything after it - the report should point at the edit, not drown it.
    Signatures are only counted when present: an unsigned chain is still
    tamper-evident, just not tamper-proof.
    """
    import hashlib

    first_break = None
    previous = ""
    signed_ok = 0
    signed_bad = 0
    for index, entry in enumerate(entries):
        expected = hashlib.sha256(
            f"{previous}|{entry.get('source_sha256', '')}"
            f"|{entry.get('output_sha256', '')}"
            f"|{entry.get('verdict', '')}".encode("ascii")
        ).hexdigest()
        if index and expected != entry.get("chain", "") and first_break is None:
            first_break = index
        signature = entry.get("signature", "")
        if signature:
            if verify_signature(entry.get("chain", ""), signature):
                signed_ok += 1
            else:
                signed_bad += 1
        previous = entry.get("chain", "")
    return {
        "count": len(entries),
        "intact": first_break is None,
        "first_break": first_break,
        "signed_ok": signed_ok,
        "signed_bad": signed_bad,
    }


def record_attestation(raw_text: str, docx_bytes: bytes, ok: bool,
                       subject: str = "") -> dict:
    """
    Compute the attestation chained onto the last one in the audit log, and
    record it. Storage failures never block clinical work - the attestation
    is still returned for the caller to show.
    """
    import json

    previous = ""
    try:
        import storage

        for event in storage.get_store().events(limit=500):
            if event.kind == "audit.attest":
                try:
                    previous = json.loads(event.detail).get("chain", "")
                except Exception:
                    previous = ""
                break
        record = attestation(raw_text, docx_bytes, ok, previous_chain=previous)
        storage.log("audit.attest", subject or "report",
                    detail=json.dumps(record, ensure_ascii=False))
        return record
    except Exception:
        return attestation(raw_text, docx_bytes, ok, previous_chain=previous)
